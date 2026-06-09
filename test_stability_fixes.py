#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
稳定性修复测试脚本 - 验证用户提出的4个问题是否已修复
"""
import sys
from pathlib import Path
from datetime import datetime
import json

sys.path.insert(0, str(Path(__file__).resolve().parent))

from core import (
    ProjectInfo,
    MaterialList,
    TemplateGenerator,
    ContentValidator,
    PackageOutput,
    RecordManager,
    SubmissionPreview
)
from config import TRADING_SCENARIOS, REQUIRED_MATERIALS, OUTPUT_DIR


def test_project_info_scene_property():
    """测试1: 验证ProjectInfo的scene属性别名，确保交易场景字段名兼容"""
    print("\n" + "="*60)
    print("测试1: ProjectInfo字段名兼容性 (scene / trading_scene)")
    print("="*60)

    pi = ProjectInfo()

    pi.scene = "金融风控"
    print(f"  设置 pi.scene = '金融风控'")
    print(f"  pi.trading_scene = '{pi.trading_scene}' (期望: 金融风控)")
    assert pi.trading_scene == "金融风控", "scene属性设置应该同步到trading_scene"

    pi.trading_scene = "市场营销"
    print(f"  设置 pi.trading_scene = '市场营销'")
    print(f"  pi.scene = '{pi.scene}' (期望: 市场营销)")
    assert pi.scene == "市场营销", "trading_scene属性设置应该同步到scene"

    pi.usage_limits = "仅限内部使用"
    print(f"  设置 pi.usage_limits = '仅限内部使用'")
    print(f"  pi.usage_restrictions = '{pi.usage_restrictions}' (期望: 仅限内部使用)")
    assert pi.usage_restrictions == "仅限内部使用", "usage_limits属性设置应该同步到usage_restrictions"

    print("  ✓ ProjectInfo字段别名测试通过!")
    return True


def test_material_list_to_dict_compatibility():
    """测试2: 验证MaterialList的to_dict方法同时支持items和materials键"""
    print("\n" + "="*60)
    print("测试2: MaterialList序列化兼容性 (items / materials)")
    print("="*60)

    ml = MaterialList()
    ml.mark_provided("CPMS", "/test/product_desc.docx")
    ml.mark_generated("SQMS", True, "batch_001")

    data = ml.to_dict()
    print(f"  to_dict() 包含 'items' 键: {'items' in data}")
    print(f"  to_dict() 包含 'materials' 键: {'materials' in data}")
    assert "items" in data, "应该包含items键"
    assert "materials" in data, "应该包含materials键"
    assert len(data["items"]) == len(data["materials"]), "items和materials应该有相同数量的元素"
    assert data["items"][0]["code"] == data["materials"][0]["code"], "items和materials应该包含相同的数据"

    ml2 = MaterialList()
    ml2.from_dict({"materials": data["materials"]})
    print(f"  from_dict() 使用 'materials' 键恢复: {len(ml2.items)} 个材料")
    assert len(ml2.items) == len(ml.items), "使用materials键应该能正确恢复"

    ml3 = MaterialList()
    ml3.from_dict({"items": data["items"]})
    print(f"  from_dict() 使用 'items' 键恢复: {len(ml3.items)} 个材料")
    assert len(ml3.items) == len(ml.items), "使用items键应该能正确恢复"

    print("  ✓ MaterialList序列化兼容性测试通过!")
    return True


def test_submission_preview_fields():
    """测试3: 验证提交预览的字段与GUI期望的字段匹配"""
    print("\n" + "="*60)
    print("测试3: SubmissionPreview字段匹配 (name, source, size, confirmed)")
    print("="*60)

    pi = ProjectInfo()
    pi.project_id = "TEST_STABLE_001"
    pi.product_code = "STABLE001"
    pi.product_name = "稳定性测试产品"
    pi.contact_person = "测试员"
    pi.contact_phone = "13800138000"
    pi.scene = TRADING_SCENARIOS[0]
    pi.valid_from = "2024-01-01"
    pi.valid_to = "2024-12-31"
    pi.data_source = "测试数据源"
    pi.update_frequency = "每日"

    ml = MaterialList()
    batch = ml.create_batch("稳定性测试批次", "overwrite", ["CPMS", "SQMS", "YLDQ", "SCRZ"])

    output_dir = OUTPUT_DIR / pi.project_id
    output_dir.mkdir(parents=True, exist_ok=True)

    test_file_cpms = output_dir / "TEST_STABLE_001_CPMS.docx"
    test_file_sqms = output_dir / "TEST_STABLE_001_SQMS.docx"
    test_file_scrz = output_dir / "TEST_STABLE_001_SCRZ.pdf"

    from docx import Document
    doc = Document()
    doc.add_heading("产品说明", level=1)
    doc.add_paragraph("这是测试文档内容")
    doc.save(str(test_file_cpms))
    doc.save(str(test_file_sqms))

    with open(test_file_scrz, "w") as f:
        f.write("测试数据来源证明")

    ml.mark_generated("CPMS", True, batch.batch_id)
    ml.mark_generated("SQMS", True, batch.batch_id)
    ml.mark_provided("SCRZ", str(test_file_scrz))

    generated_files = {
        "CPMS": str(test_file_cpms),
        "SQMS": str(test_file_sqms),
    }

    ml.confirm_batch_files(batch.batch_id, [str(test_file_cpms), str(test_file_sqms), str(test_file_scrz)])

    cv = ContentValidator(pi, ml)
    vr = cv.validate_all()

    packager = PackageOutput(pi, ml, generated_files, vr)
    preview = packager.get_submission_preview()

    print(f"  待打包文件数量: {len(preview.files_to_include)}")
    for i, f in enumerate(preview.files_to_include, 1):
        print(f"\n  文件 {i}:")
        print(f"    - 'name' 字段: {f.get('name', 'MISSING!')}")
        print(f"    - 'source' 字段: {f.get('source', 'MISSING!')}")
        print(f"    - 'size' 字段: {f.get('size', 'MISSING!')}")
        print(f"    - 'confirmed' 字段: {f.get('confirmed', 'MISSING!')}")
        print(f"    - 'original_path' 字段: {f.get('original_path', 'MISSING!')}")

        assert "name" in f, "应该包含name字段"
        assert "source" in f, "应该包含source字段"
        assert "size" in f, "应该包含size字段"
        assert "confirmed" in f, "应该包含confirmed字段"
        assert "original_path" in f, "应该包含original_path字段"

    print(f"\n  缺失必填材料数量: {len(preview.missing_required)}")
    print(f"  缺失选填材料数量: {len(preview.missing_optional)}")
    print(f"  敏感词警告数量: {len(preview.sensitive_warnings)}")
    print(f"  校验错误数量: {len(preview.validation_errors)}")
    print(f"  校验警告数量: {len(preview.validation_warnings)}")
    print(f"  可以提交: {preview.can_submit}")

    for i, w in enumerate(preview.sensitive_warnings, 1):
        print(f"\n  敏感词警告 {i}:")
        print(f"    - 'file' 字段: {w.get('file', 'MISSING!')}")
        print(f"    - 'words' 字段: {w.get('words', 'MISSING!')}")
        print(f"    - 'line' 字段: {w.get('line', 'MISSING!')}")

        assert "file" in w, "敏感词警告应该包含file字段"
        assert "words" in w, "敏感词警告应该包含words字段"
        assert "line" in w, "敏感词警告应该包含line字段"

    print("  ✓ SubmissionPreview字段匹配测试通过!")

    for f in [test_file_cpms, test_file_sqms, test_file_scrz]:
        if f.exists():
            f.unlink()

    return True


def test_build_package_result_fields():
    """测试4: 验证build_package返回的字段与GUI期望匹配"""
    print("\n" + "="*60)
    print("测试4: build_package返回字段匹配 (file_count, zip_path, checklist_path)")
    print("="*60)

    pi = ProjectInfo()
    pi.project_id = "TEST_PKG_001"
    pi.product_code = "PKG001"
    pi.product_name = "打包测试产品"
    pi.contact_person = "测试员"
    pi.contact_phone = "13800138000"
    pi.scene = TRADING_SCENARIOS[0]
    pi.valid_from = "2024-01-01"
    pi.valid_to = "2024-12-31"
    pi.data_source = "测试数据源"
    pi.update_frequency = "每日"

    ml = MaterialList()
    batch = ml.create_batch("打包测试批次", "overwrite", ["CPMS", "SQMS", "YLDQ"])

    output_dir = OUTPUT_DIR / pi.project_id
    output_dir.mkdir(parents=True, exist_ok=True)

    test_file_cpms = output_dir / "TEST_PKG_001_CPMS.docx"
    test_file_sqms = output_dir / "TEST_PKG_001_SQMS.docx"
    test_file_yldq = output_dir / "TEST_PKG_001_YLDQ.xlsx"

    from docx import Document
    doc = Document()
    doc.add_heading("产品说明", level=1)
    doc.add_paragraph("这是测试文档内容")
    doc.save(str(test_file_cpms))
    doc.save(str(test_file_sqms))

    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["字段名", "类型", "描述"])
    ws.append(["id", "string", "唯一标识"])
    wb.save(str(test_file_yldq))

    ml.mark_generated("CPMS", True, batch.batch_id)
    ml.mark_generated("SQMS", True, batch.batch_id)
    ml.mark_generated("YLDQ", True, batch.batch_id)

    generated_files = {
        "CPMS": str(test_file_cpms),
        "SQMS": str(test_file_sqms),
        "YLDQ": str(test_file_yldq),
    }

    ml.confirm_batch_files(batch.batch_id, [str(test_file_cpms), str(test_file_sqms), str(test_file_yldq)])

    cv = ContentValidator(pi, ml)
    vr = cv.validate_all()

    packager = PackageOutput(pi, ml, generated_files, vr)
    result = packager.build_package(ignore_errors=True)

    print(f"  打包成功: {result.get('success', False)}")
    print(f"  'zip_path' 字段: {result.get('zip_path', 'MISSING!')}")
    print(f"  'checklist_path' 字段: {result.get('checklist_path', 'MISSING!')}")
    print(f"  'file_count' 字段: {result.get('file_count', 'MISSING!')}")
    print(f"  'output_files' 字段: {len(result.get('output_files', []))} 个文件")

    assert "success" in result, "应该包含success字段"
    assert result["success"] == True, "打包应该成功"
    assert "zip_path" in result and result["zip_path"], "应该包含zip_path字段"
    assert "checklist_path" in result and result["checklist_path"], "应该包含checklist_path字段"
    assert "file_count" in result, "应该包含file_count字段"
    assert "output_files" in result, "应该包含output_files字段"

    assert Path(result["zip_path"]).exists(), "压缩包文件应该存在"
    assert Path(result["checklist_path"]).exists(), "清单文件应该存在"
    assert result["file_count"] == len(result["output_files"]), "file_count应该等于output_files的长度"

    print("  ✓ build_package返回字段匹配测试通过!")

    for f in [test_file_cpms, test_file_sqms, test_file_yldq]:
        if f.exists():
            f.unlink()
    if Path(result["zip_path"]).exists():
        Path(result["zip_path"]).unlink()
    if Path(result["checklist_path"]).exists():
        Path(result["checklist_path"]).unlink()

    return True


def test_record_save_and_restore():
    """测试5: 验证记录保存和恢复功能，确保压缩包路径等信息能正确保存和恢复"""
    print("\n" + "="*60)
    print("测试5: 记录保存和恢复 (压缩包路径、材料状态等)")
    print("="*60)

    rm = RecordManager()

    pi = ProjectInfo()
    pi.project_id = "REC_TEST_001"
    pi.product_code = "REC001"
    pi.product_name = "记录测试产品"
    pi.contact_person = "测试员"
    pi.scene = TRADING_SCENARIOS[0]

    ml = MaterialList()
    batch = ml.create_batch("记录测试批次", "overwrite")
    ml.mark_generated("CPMS", True, batch.batch_id)
    ml.mark_provided("SCRZ", "/test/source.pdf")
    ml.confirm_batch_files(batch.batch_id, ["/test/product_desc.docx", "/test/source.pdf"])

    generated_files = {
        "CPMS": "/output/REC_TEST_001_CPMS.docx",
        "SQMS": "/output/REC_TEST_001_SQMS.docx",
    }

    validation_result = {
        "passed": True,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "errors": [],
        "warnings": [],
        "sensitive_warnings": []
    }

    zip_path = "/output/REC_TEST_001_20240101.zip"
    checklist_path = "/output/REC_TEST_001_checklist.xlsx"

    record_data = {
        "project_id": pi.project_id,
        "product_code": pi.product_code,
        "product_name": pi.product_name,
        "contact_person": pi.contact_person,
        "scene": pi.scene,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "batch_id": batch.batch_id,
        "batch_name": batch.batch_name,
        "project_info": pi.to_dict(),
        "materials": ml.to_dict(),
        "generated_files": generated_files,
        "validation_result": validation_result,
        "zip_path": zip_path,
        "checklist_path": checklist_path,
        "file_count": 5
    }

    saved_record = rm.add_record(record_data)
    print(f"  记录已保存, project_id: {saved_record['project_id']}")
    print(f"  保存的 zip_path: {saved_record.get('zip_path', 'MISSING!')}")
    print(f"  保存的 checklist_path: {saved_record.get('checklist_path', 'MISSING!')}")
    print(f"  保存的 file_count: {saved_record.get('file_count', 'MISSING!')}")

    assert "zip_path" in saved_record, "记录应该包含zip_path"
    assert saved_record["zip_path"] == zip_path, "zip_path应该正确保存"
    assert "checklist_path" in saved_record, "记录应该包含checklist_path"
    assert saved_record["checklist_path"] == checklist_path, "checklist_path应该正确保存"

    search_results = rm.search_records(project_id=pi.project_id)
    assert len(search_results) == 1, "应该能搜索到刚才保存的记录"

    found_record = search_results[0]
    print(f"\n  搜索到记录, zip_path: {found_record.get('zip_path', 'MISSING!')}")
    assert found_record["zip_path"] == zip_path, "搜索到的记录应该包含正确的zip_path"

    new_pi = ProjectInfo()
    new_ml = MaterialList()
    new_gf = {}
    new_vr = None

    restore_success = rm.restore_record(found_record, new_pi, new_ml, new_gf, new_vr)
    print(f"\n  恢复记录成功: {restore_success}")
    print(f"  恢复的 project_id: {new_pi.project_id}")
    print(f"  恢复的 product_name: {new_pi.product_name}")
    print(f"  恢复的 scene: {new_pi.scene}")
    print(f"  恢复的 generated_files: {list(new_gf.keys())}")

    assert restore_success, "恢复记录应该成功"
    assert new_pi.project_id == pi.project_id, "恢复的project_id应该正确"
    assert new_pi.product_name == pi.product_name, "恢复的product_name应该正确"
    assert new_pi.scene == pi.scene, "恢复的scene应该正确"
    assert "CPMS" in new_gf, "恢复的generated_files应该包含CPMS"

    cpms_item = new_ml.get_item("CPMS")
    scrz_item = new_ml.get_item("SCRZ")
    print(f"\n  恢复的材料状态:")
    print(f"    CPMS - generated: {cpms_item.generated}, file_path: {cpms_item.file_path}")
    print(f"    SCRZ - provided: {scrz_item.provided}, file_path: {scrz_item.file_path}")

    assert cpms_item.generated == True, "CPMS应该标记为已生成"
    assert scrz_item.provided == True, "SCRZ应该标记为已提供"
    assert scrz_item.file_path == "/test/source.pdf", "SCRZ的文件路径应该正确"

    rm.delete_record(pi.project_id, batch.batch_id, saved_record["created_at"])
    remaining = rm.search_records(project_id=pi.project_id)
    assert len(remaining) == 0, "测试记录应该被删除"

    print("  ✓ 记录保存和恢复测试通过!")
    return True


def test_material_status_display():
    """测试6: 验证材料状态显示逻辑 (已生成/已上传/未上传)"""
    print("\n" + "="*60)
    print("测试6: 材料状态显示逻辑 (已生成/已上传/未上传)")
    print("="*60)

    ml = MaterialList()

    for item in ml.items:
        if item.generated:
            status = "已生成"
        elif item.file_path:
            status = "已上传"
        else:
            status = "未上传"
        print(f"  {item.code} - {item.name}: {status}")

    ml.mark_generated("CPMS", True, "batch_001")
    ml.mark_provided("SCRZ", "/test/source.pdf")

    print("\n  更新后:")
    materials_data = ml.to_dict()
    materials = materials_data.get("materials", materials_data.get("items", []))
    for m in materials:
        if m.get("generated"):
            status = "已生成"
        elif m.get("file_path"):
            status = "已上传"
        else:
            status = "未上传"
        print(f"  {m['code']} - {m['name']}: {status}")

        if m["code"] == "CPMS":
            assert status == "已生成", "CPMS应该是已生成状态"
        elif m["code"] == "SCRZ":
            assert status == "已上传", "SCRZ应该是已上传状态"
        elif m["code"] == "SQMS":
            assert status == "未上传", "SQMS应该是未上传状态"

    print("  ✓ 材料状态显示逻辑测试通过!")
    return True


def test_trading_scenarios_constant():
    """测试7: 验证TRADING_SCENARIOS常量可以正常使用"""
    print("\n" + "="*60)
    print("测试7: TRADING_SCENARIOS常量可用性")
    print("="*60)

    from config import TRADING_SCENARIOS as SCENES

    print(f"  交易场景数量: {len(SCENES)}")
    for i, scene in enumerate(SCENES, 1):
        print(f"    {i}. {scene}")

    assert len(SCENES) >= 8, "应该至少有8个交易场景"
    assert "金融风控" in SCENES, "应该包含金融风控场景"
    assert "市场营销" in SCENES, "应该包含市场营销场景"

    print("  ✓ TRADING_SCENARIOS常量测试通过!")
    return True


def main():
    print("\n" + "="*60)
    print("数据要素交易材料生成工具 - 稳定性修复验证测试")
    print("="*60)

    tests = [
        test_project_info_scene_property,
        test_material_list_to_dict_compatibility,
        test_submission_preview_fields,
        test_build_package_result_fields,
        test_record_save_and_restore,
        test_material_status_display,
        test_trading_scenarios_constant
    ]

    passed = 0
    failed = 0

    for test in tests:
        try:
            if test():
                passed += 1
            else:
                failed += 1
        except Exception as e:
            print(f"  ✗ 测试失败: {str(e)}")
            import traceback
            traceback.print_exc()
            failed += 1

    print("\n" + "="*60)
    print(f"测试结果: {passed} 通过, {failed} 失败")
    print("="*60)

    if failed == 0:
        print("\n✓ 所有稳定性修复测试通过!")
        print("\n修复的问题总结:")
        print("  1. ✓ ProjectInfo添加了scene和usage_limits属性别名，确保字段名兼容")
        print("  2. ✓ MaterialList的to_dict()同时返回items和materials键，确保兼容性")
        print("  3. ✓ SubmissionPreview的字段与GUI期望匹配 (name, source, size, confirmed)")
        print("  4. ✓ build_package返回file_count字段，与GUI期望匹配")
        print("  5. ✓ 记录保存包含zip_path, checklist_path, file_count等完整信息")
        print("  6. ✓ 材料状态区分已生成/已上传/未上传，显示更精确")
        print("  7. ✓ TRADING_SCENARIOS常量可正常导入和使用")
        return 0
    else:
        print("\n✗ 部分测试失败，请检查代码!")
        return 1


if __name__ == "__main__":
    sys.exit(main())
