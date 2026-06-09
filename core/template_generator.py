from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from core.project_info import ProjectInfo
from core.material_list import MaterialList
from config import OUTPUT_DIR, TEMPLATES_DIR


@dataclass
class GenerationResult:
    material_code: str
    material_name: str
    file_path: str
    success: bool
    error_message: str = ""
    is_new: bool = False
    batch_id: str = ""


GENERATE_STRATEGIES = ["overwrite", "new_version", "update_selected"]


class TemplateGenerator:
    GENERATE_STRATEGIES = GENERATE_STRATEGIES

    def __init__(self, project_info: ProjectInfo, material_list: MaterialList = None):
        self.project = project_info
        self.materials = material_list
        self.output_dir = OUTPUT_DIR / project_info.project_id
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._existing_files = self._scan_existing_files()

    def _scan_existing_files(self) -> Dict[str, str]:
        existing = {}
        if not self.output_dir.exists():
            return existing
        for f in self.output_dir.iterdir():
            if f.is_file() and f.suffix in ['.docx', '.xlsx']:
                if '_CPMS_' in f.name or 'CPMS' in f.stem:
                    existing['CPMS'] = str(f)
                elif '_SQMS_' in f.name or 'SQMS' in f.stem:
                    existing['SQMS'] = str(f)
                elif '_YLDQ_' in f.name or 'YLDQ' in f.stem:
                    existing['YLDQ'] = str(f)
        return existing

    def generate_all(self, strategy: str = "overwrite",
                     selected_codes: List[str] = None,
                     batch_name: str = "") -> Dict:
        if strategy not in self.GENERATE_STRATEGIES:
            strategy = "overwrite"

        material_codes = selected_codes or ["CPMS", "SQMS", "YLDQ"]
        results = {}
        generation_results = []

        batch = None
        if self.materials:
            batch = self.materials.create_batch(
                name=batch_name,
                strategy=strategy,
                material_codes=material_codes
            )

        for code in material_codes:
            if code == "CPMS" and (selected_codes is None or "CPMS" in selected_codes):
                result = self._generate_with_strategy(
                    "product_desc", code, "产品说明文档",
                    self.generate_product_description, strategy, batch
                )
                generation_results.append(result)
                if result.success:
                    results["product_desc"] = result.file_path
                elif result.error_message:
                    results["product_desc_error"] = result.error_message

            elif code == "SQMS" and (selected_codes is None or "SQMS" in selected_codes):
                result = self._generate_with_strategy(
                    "auth_desc", code, "授权说明文档",
                    self.generate_authorization_description, strategy, batch
                )
                generation_results.append(result)
                if result.success:
                    results["auth_desc"] = result.file_path
                elif result.error_message:
                    results["auth_desc_error"] = result.error_message

            elif code == "YLDQ" and (selected_codes is None or "YLDQ" in selected_codes):
                result = self._generate_with_strategy(
                    "sample_fields", code, "样例字段清单",
                    self.generate_sample_fields_list, strategy, batch
                )
                generation_results.append(result)
                if result.success:
                    results["sample_fields"] = result.file_path
                elif result.error_message:
                    results["sample_fields_error"] = result.error_message

        results["_generation_results"] = [r.__dict__ for r in generation_results]
        results["_batch_id"] = batch.batch_id if batch else ""
        results["_strategy"] = strategy

        if self.materials and batch:
            confirmed = [r.file_path for r in generation_results if r.success]
            for item in self.materials.get_provided():
                if item.file_path and not item.generated:
                    confirmed.append(item.file_path)
            self.materials.confirm_batch_files(batch.batch_id, confirmed)

        return results

    def _generate_with_strategy(self, key: str, code: str, name: str,
                                generate_func, strategy: str, batch) -> GenerationResult:
        batch_id = batch.batch_id if batch else ""
        existing = self._existing_files.get(code)

        try:
            if strategy == "overwrite":
                file_path = generate_func()
                is_new = not existing or existing != file_path
                if self.materials:
                    self.materials.mark_generated(code, True, batch_id)
                return GenerationResult(code, name, file_path, True, "", is_new, batch_id)

            elif strategy == "new_version":
                file_path = generate_func(new_version=True)
                if self.materials:
                    self.materials.mark_generated(code, True, batch_id)
                return GenerationResult(code, name, file_path, True, "", True, batch_id)

            elif strategy == "update_selected":
                file_path = generate_func()
                if self.materials:
                    self.materials.mark_generated(code, True, batch_id)
                return GenerationResult(code, name, file_path, True, "", not existing, batch_id)

        except Exception as e:
            return GenerationResult(code, name, "", False, str(e), False, batch_id)

    def generate_product_description(self, new_version: bool = False) -> str:
        doc = Document()
        title = doc.add_heading(f"{self.project.product_name} - 产品说明", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        self._add_section(doc, "一、产品概述", [
            f"产品名称：{self.project.product_name}",
            f"产品编码：{self.project.product_code}",
            f"交易场景：{self.project.trading_scene}",
            f"数据来源：{self.project.data_source}",
            f"更新频率：{self.project.update_frequency}",
            f"数据量级：{self.project.data_volume}"
        ])
        self._add_section(doc, "二、产品描述", [self.project.data_description or "暂无详细描述。"])
        self._add_section(doc, "三、数据字段说明", [
            "本产品包含以下核心字段（完整清单参见《样例字段清单》）："
        ])
        if self.project.sample_fields:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["字段名称", "字段类型", "描述", "样例值"]
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                for p in hdr_cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
            for field in self.project.sample_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = str(field.get("field_name", ""))
                row_cells[1].text = str(field.get("field_type", ""))
                row_cells[2].text = str(field.get("description", ""))
                row_cells[3].text = str(field.get("sample_value", ""))
            doc.add_paragraph()
        else:
            doc.add_paragraph("暂无字段信息。")
            doc.add_paragraph()
        self._add_section(doc, "四、使用限制", [
            self.project.usage_restrictions or "本产品的使用需遵守相关法律法规及平台规定。"
        ])
        self._add_section(doc, "五、联系方式", [
            f"联系人：{self.project.contact_person or '待定'}",
            f"联系电话：{self.project.contact_phone or '待定'}",
            f"电子邮箱：{self.project.contact_email or '待定'}"
        ])
        self._add_footer(doc)
        filename = self._get_filename("CPMS_产品说明", ".docx", new_version)
        filepath = self.output_dir / filename
        doc.save(filepath)
        return str(filepath)

    def generate_authorization_description(self, new_version: bool = False) -> str:
        doc = Document()
        title = doc.add_heading(f"{self.project.product_name} - 授权说明", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        self._add_section(doc, "一、授权方信息", [
            f"数据来源方：{self.project.data_source}",
            f"联系人：{self.project.contact_person or '待定'}",
            f"联系电话：{self.project.contact_phone or '待定'}"
        ])
        self._add_section(doc, "二、授权内容", [
            f"产品名称：{self.project.product_name}",
            f"授权交易场景：{self.project.trading_scene}",
            "授权方式：非独占性授权，被授权方不得转授权。"
        ])
        self._add_section(doc, "三、授权期限", [
            f"有效期自：{self.project.valid_from or '待定'}",
            f"有效期至：{self.project.valid_to or '待定'}"
        ])
        self._add_section(doc, "四、使用限制", [
            self.project.usage_restrictions or "使用方应遵守数据安全相关法律法规，不得用于授权范围以外的用途。",
            "使用方应采取必要的安全措施保护数据安全，防止数据泄露。"
        ])
        self._add_section(doc, "五、权利声明", [
            "授权方保证对所提供的数据产品拥有合法的权利。",
            "数据产品的使用不得侵犯任何第三方的合法权益。"
        ])
        doc.add_paragraph()
        doc.add_paragraph("授权方（盖章）：")
        doc.add_paragraph()
        doc.add_paragraph("日期：")
        self._add_footer(doc)
        filename = self._get_filename("SQMS_授权说明", ".docx", new_version)
        filepath = self.output_dir / filename
        doc.save(filepath)
        return str(filepath)

    def generate_sample_fields_list(self, new_version: bool = False) -> str:
        wb = Workbook()
        ws = wb.active
        ws.title = "样例字段清单"
        headers = ["序号", "字段名称", "字段类型", "字段描述", "样例值", "备注"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        info_rows = [
            ["产品名称", self.project.product_name],
            ["产品编码", self.project.product_code],
            ["数据来源", self.project.data_source],
            ["更新频率", self.project.update_frequency],
            ["导出日期", datetime.now().strftime("%Y-%m-%d")]
        ]
        for row_idx, (key, value) in enumerate(info_rows, 2):
            ws.cell(row=row_idx, column=1, value=key).font = Font(bold=True)
            ws.cell(row=row_idx, column=2, value=value)
        start_row = len(info_rows) + 3
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4472C4", end_color="#4472C4", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        if self.project.sample_fields:
            for idx, field in enumerate(self.project.sample_fields, 1):
                row = start_row + idx
                ws.cell(row=row, column=1, value=idx).alignment = Alignment(horizontal="center")
                ws.cell(row=row, column=2, value=field.get("field_name", ""))
                ws.cell(row=row, column=3, value=field.get("field_type", ""))
                ws.cell(row=row, column=4, value=field.get("description", ""))
                ws.cell(row=row, column=5, value=field.get("sample_value", ""))
                ws.cell(row=row, column=6, value="")
                for col in range(1, 7):
                    ws.cell(row=row, column=col).border = thin_border
        else:
            row = start_row + 1
            ws.cell(row=row, column=1, value="暂无字段信息").alignment = Alignment(horizontal="center")
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        for col in range(1, 7):
            ws.column_dimensions[chr(64 + col)].width = 20
        filename = self._get_filename("YLDQ_样例字段清单", ".xlsx", new_version)
        filepath = self.output_dir / filename
        wb.save(filepath)
        return str(filepath)

    def _get_filename(self, suffix: str, extension: str, new_version: bool = False) -> str:
        date_str = datetime.now().strftime("%Y%m%d")
        time_str = datetime.now().strftime("%H%M%S")
        scene = self._safe_filename(self.project.trading_scene[:4] if self.project.trading_scene else "场景")
        code = self._safe_filename(self.project.product_code if self.project.product_code else self.project.project_id)
        if new_version:
            return f"{code}_{scene}_{date_str}_{time_str}_{suffix}{extension}"
        return f"{code}_{scene}_{date_str}_{suffix}{extension}"

    def _safe_filename(self, filename: str) -> str:
        invalid_chars = '<>:"/\\|?*'
        for ch in invalid_chars:
            filename = filename.replace(ch, '_')
        return filename.strip() or "unnamed"

    def _add_section(self, doc: Document, title: str, contents: List[str]) -> None:
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)
        for content in contents:
            p = doc.add_paragraph(content)
            for run in p.runs:
                run.font.size = Pt(11)
        doc.add_paragraph()

    def _add_footer(self, doc: Document) -> None:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 文档编号：{self.project.project_id}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
